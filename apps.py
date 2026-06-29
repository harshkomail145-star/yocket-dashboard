import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime
import io
from docx import Document
from docx.shared import Inches

# --- 1. PAGE CONFIG & THEME ---
st.set_page_config(page_title="Fall 26 Analytics", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    .stMetric { background-color: #ffffff; padding: 20px; border-radius: 12px; border: 1px solid #e2e8f0; border-top: 4px solid #4f46e5; box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1); }
    .section-header { background-color: #ffffff; padding: 15px; border-radius: 8px; border-left: 5px solid #4f46e5; margin-top: 30px; margin-bottom: 15px; box-shadow: 0 1px 2px rgba(0,0,0,0.05);}
    </style>
    """, unsafe_allow_html=True)

st.title("📊 Fall 26 Command Center")

with st.sidebar:
    st.header("⚙️ Global Filters")
    selected_banks = st.multiselect("Select Bank Partners", ["Auxilo", "Credila", "Avanse", "SBI", "ICICI"], default=["Credila"])
    st.divider()
    st.caption("UI Mode: MOCK DATA 🟢")

# ==========================================
# PRE-BUILD ALL CHARTS FOR THE EXPORT
# ==========================================
# 1. Top Metrics
stages = ['Shared', 'Login', 'Sanction', 'PF']
fall_25_data = [2067, 1752, 908, 467]
fall_26_data = [2855, 2225, 1110, 588]
fig_top_metrics = go.Figure()
fig_top_metrics.add_trace(go.Bar(name='Fall 25', x=stages, y=fall_25_data, marker_color='#6a96b9', text=fall_25_data, textposition='outside'))
fig_top_metrics.add_trace(go.Bar(name='Fall 26', x=stages, y=fall_26_data, marker_color='#1f4e71', text=fall_26_data, textposition='outside'))
fig_top_metrics.update_layout(barmode='group', plot_bgcolor='rgba(0,0,0,0)')

# 2. YoY Monthly
mock_yoy_monthly = pd.DataFrame({'Month': ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun'], 'Fall 26': [219, 397, 444, 436, 377, 352], 'Fall 25': [243, 297, 380, 322, 294, 231]})
fig_yoy_bar = go.Figure(data=[
    go.Bar(name='Fall 26', x=mock_yoy_monthly['Month'], y=mock_yoy_monthly['Fall 26'], marker_color='#60a5fa'),
    go.Bar(name='Fall 25', x=mock_yoy_monthly['Month'], y=mock_yoy_monthly['Fall 25'], marker_color='#ef4444')
])
fig_yoy_bar.update_layout(barmode='group', plot_bgcolor='rgba(0,0,0,0)')

# 3. Funnel
totals = [2783, 2148, 1021, 504]
fig_funnel = go.Figure(go.Funnel(y=stages, x=totals, textinfo="value+percent previous"))
fig_funnel.update_layout(plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)')

# 4. Lost Potential (Section 5)
fig_dual = make_subplots(specs=[[{"secondary_y": True}]])
stages_lost = ['Lost from BP', 'Lost from Login', 'Lost from Sanction']
fig_dual.add_trace(go.Bar(name='Total Files', x=stages_lost, y=[200, 300, 100], marker_color='#60a5fa'), secondary_y=False)
fig_dual.add_trace(go.Bar(name='Went Ahead', x=stages_lost, y=[145, 204, 85], marker_color='#ef4444'), secondary_y=False)
fig_dual.add_trace(go.Scatter(name='% Lost Potential', x=stages_lost, y=[72.5, 68.0, 85.0], mode='lines+markers', line=dict(width=4, color='#fbbf24')), secondary_y=True)
fig_dual.update_layout(barmode='group', plot_bgcolor='rgba(0,0,0,0)')

# ==========================================
# DOCX GENERATOR FUNCTION
# ==========================================
def generate_docx_report():
    doc = Document()
    doc.add_heading(f'Fall 26 Command Center - Report ({datetime.now().strftime("%Y-%m-%d")})', 0)
    
    # Helper to convert plotly chart to image and add to doc
    def add_plot_to_doc(fig, title):
        doc.add_heading(title, level=2)
        img_bytes = fig.to_image(format="png", width=700, height=400)
        doc.add_picture(io.BytesIO(img_bytes), width=Inches(6.0))
    
    add_plot_to_doc(fig_top_metrics, '1. Y-o-Y Metric Comparison')
    add_plot_to_doc(fig_yoy_bar, '2. YoY Monthly Logins')
    add_plot_to_doc(fig_funnel, '3. Shared Lead Funnel')
    add_plot_to_doc(fig_dual, '4. Lost Potential Analysis (Section 5)')
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- THE NEW DOWNLOAD BUTTON ---
st.download_button(
    label="📥 Download FULL Visual Report (.docx)",
    data=generate_docx_report(),
    file_name=f"Fall_26_Visual_Report.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    help="Download this file and open it with Google Docs to see all charts."
)

st.divider()

# ==========================================
# UI RENDERING (What the user actually sees on the webpage)
# ==========================================
st.markdown('<div class="section-header"><h2>📈 1. Y-o-Y Metrics</h2></div>', unsafe_allow_html=True)
st.plotly_chart(fig_top_metrics, use_container_width=True)

st.markdown('<div class="section-header"><h2>📅 2. M-o-M Logins</h2></div>', unsafe_allow_html=True)
st.plotly_chart(fig_yoy_bar, use_container_width=True)

st.markdown('<div class="section-header"><h2>🧬 3. Shared Leads Funnel</h2></div>', unsafe_allow_html=True)
st.plotly_chart(fig_funnel, use_container_width=True)

st.markdown('<div class="section-header"><h2>🚨 5. Lost Analysis (Now Fully Included!)</h2></div>', unsafe_allow_html=True)
st.plotly_chart(fig_dual, use_container_width=True)
